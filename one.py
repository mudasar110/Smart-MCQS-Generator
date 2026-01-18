import streamlit as st
import re
import random
import string
import io
import sys

# --- SAFE IMPORTS (Deployment Ready) ---
# If libraries are missing, the app will disable those features instead of crashing
PDF_ENABLED = False
try:
    from fpdf import FPDF
    PDF_ENABLED = True
except ImportError:
    pass

DOCX_ENABLED = False
try:
    from docx import Document
    DOCX_ENABLED = True
except ImportError:
    pass

pypdf = None
try:
    import pypdf
    # Try to import specific errors if available for better handling
    try:
        from pypdf.errors import PdfReadError, DependencyError
    except ImportError:
        # Older versions might not have these classes in errors
        PdfReadError = Exception 
        DependencyError = Exception
except ImportError:
    pass

# --- CONFIGURATION ---
st.set_page_config(
    page_title="SP BROTHERS - Smart Exam",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
if 'generated_mcqs' not in st.session_state:
    st.session_state.generated_mcqs = []

# --- SP BROTHERS THEME STYLING ---
st.markdown("""
    <style>
    :root {
        --primary: #00e5ff;
        --primary-dim: #00e5ff33;
        --secondary: #2979ff;
        --bg-dark: #0a0e17;
        --bg-panel: #111625;
        --text-main: #e0e6ed;
        --text-muted: #94a3b8;
        --success: #00ff9d;
        --danger: #ff4757;
        --warning: #ffa502;
        --font-tech: 'Courier New', Courier, monospace;
        --font-ui: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
    }

    body {
        background-color: var(--bg-dark);
        color: var(--text-main);
        font-family: var(--font-ui);
    }

    .stApp {
        background-color: var(--bg-dark);
    }

    .main-header {
        font-family: var(--font-ui);
        font-size: 2.2rem;
        font-weight: 800;
        text-align: center;
        margin-bottom: 2rem;
        padding-bottom: 1rem;
        border-bottom: 1px solid var(--primary-dim);
        background: linear-gradient(45deg, var(--primary), var(--secondary));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-transform: uppercase;
        letter-spacing: 2px;
    }

    .sub-header {
        text-align: center;
        color: var(--text-muted);
        font-family: var(--font-tech);
        font-size: 0.9rem;
        margin-top: -1.5rem;
        margin-bottom: 2rem;
    }

    .stTextInput > div > div > input,
    .stNumberInput > div > div > input,
    .stTextArea > div > div > textarea {
        background-color: var(--bg-panel) !important;
        color: var(--text-main) !important;
        border: 1px solid var(--secondary) !important;
        border-radius: 6px;
    }
    
    [data-testid="stFileUploader"] {
        background-color: var(--bg-panel) !important;
        border: 2px dashed var(--secondary) !important;
        border-radius: 12px;
        padding: 2rem;
        text-align: center;
    }
    
    [data-testid="stFileUploader"] p {
        color: var(--text-muted);
    }

    .stButton > button {
        background-color: var(--secondary);
        color: white;
        border: none;
        padding: 0.8rem 2rem;
        border-radius: 6px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.2s;
        text-transform: uppercase;
        letter-spacing: 1px;
        font-family: var(--font-ui);
        width: 100%;
        box-shadow: 0 4px 6px rgba(0,0,0,0.3);
    }

    .stButton > button:hover {
        background-color: var(--primary);
        color: var(--bg-dark);
        box-shadow: 0 0 15px var(--primary-dim);
        transform: translateY(-2px);
    }

    .stDownloadButton > button {
        background-color: #1e293b;
        color: var(--text-main);
        border: 1px solid var(--secondary);
        border-radius: 6px;
        font-weight: 500;
        width: 100%;
    }
    
    .stDownloadButton > button:hover {
        background-color: var(--secondary);
        color: white;
        border-color: var(--secondary);
    }

    .mcq-container {
        background-color: var(--bg-panel);
        border: 1px solid #1e293b;
        border-left: 4px solid var(--secondary);
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.2);
        transition: transform 0.2s;
    }

    .mcq-container:hover {
        border-left-color: var(--primary);
        transform: translateX(5px);
    }

    .question {
        font-size: 1.15rem;
        font-weight: bold;
        margin-bottom: 15px;
        color: var(--text-main);
    }

    .options {
        margin-left: 15px;
        color: var(--text-muted);
        line-height: 1.8;
        font-family: var(--font-tech);
    }

    .answer-badge {
        display: inline-block;
        background: rgba(0, 255, 157, 0.15);
        color: var(--success);
        border: 1px solid var(--success);
        padding: 6px 16px;
        border-radius: 50px;
        font-size: 0.9rem;
        font-weight: bold;
        margin-top: 15px;
        box-shadow: 0 0 10px rgba(0, 255, 157, 0.2);
        font-family: var(--font-tech);
    }

    h3 {
        color: var(--primary) !important;
        font-family: var(--font-tech);
        border-bottom: 1px solid var(--primary-dim);
        padding-bottom: 10px;
        margin-top: 2rem !important;
    }

    .stAlert {
        background-color: var(--bg-panel) !important;
        border: 1px solid var(--secondary) !important;
        color: var(--text-main) !important;
        border-radius: 8px;
        font-family: var(--font-tech);
    }

    .footer-custom {
        margin-top: 4rem;
        padding: 2rem;
        text-align: center;
        border-top: 1px solid #1e293b;
        color: var(--text-muted);
        font-size: 0.85rem;
    }
    
    .highlight {
        color: var(--primary);
        font-weight: bold;
    }
    
    ::-webkit-scrollbar { width: 10px; }
    ::-webkit-scrollbar-track { background: var(--bg-dark); }
    ::-webkit-scrollbar-thumb { background: var(--secondary); border-radius: 5px; }
    </style>
""", unsafe_allow_html=True)

# --- HELPER: Safe Logging ---
def safe_log(message):
    # In a real deployment, this could go to a file, but for Streamlit stderr is fine
    print(f"[SYSTEM LOG] {message}", file=sys.stderr)

# --- SMART LOGIC ENGINE ---

def clean_text(text):
    if not text: return ""
    # Strip non-ascii characters that might break regex or PDF
    text = text.encode("ascii", "ignore").decode("ascii") 
    text = re.sub(r'\n+', ' ', text)
    text = re.sub(r'[^\w\s.,!?;-]', '', text)
    return text.strip()

def get_sentences(text):
    try:
        # Split by punctuation
        sentences = re.split(r'(?<=[.!?]) +', text)
        valid = []
        for s in sentences:
            s = s.strip()
            if 20 < len(s) < 400:
                valid.append(s)
        return valid
    except Exception as e:
        safe_log(f"Sentence split error: {e}")
        return []

def extract_key_terms(sentences):
    try:
        terms = set()
        for sent in sentences:
            words = sent.split()
            if words:
                term = words[0].strip(string.punctuation)
                if term and term[0].isupper() and len(term) > 3:
                    terms.add(term)
        return list(terms)
    except Exception as e:
        safe_log(f"Key term extraction error: {e}")
        return []

def is_important_sentence(sent):
    try:
        lower = sent.lower()
        if any(p in lower for p in [" he ", " she ", " they ", " i ", " we ", " my ", " his ", " her "]):
            return False
        
        keywords = [" is ", " are ", " was ", " means ", " refers to ", 
                    " uses ", " used ", " makes ", " made ", " has ", " have ", " allows ", " helps "]
        return any(k in sent for k in keywords)
    except Exception:
        return False

def generate_smart_mcqs(text, num_questions):
    try:
        # Memory Protection: Limit text length
        MAX_TEXT_LENGTH = 50000
        if len(text) > MAX_TEXT_LENGTH:
            safe_log(f"Text truncated from {len(text)} to {MAX_TEXT_LENGTH}")
            text = text[:MAX_TEXT_LENGTH]

        text = clean_text(text)
        if not text: return []
        
        sentences = get_sentences(text)
        important_sentences = [s for s in sentences if is_important_sentence(s)]
        
        if not important_sentences:
            return []

        random.shuffle(important_sentences)
        smart_distractors = extract_key_terms(sentences)
        
        mcqs = []
        count = 0
        
        # Safety loop with iteration limit to prevent infinite loops
        max_iterations = min(len(important_sentences), num_questions * 2) 
        
        for sent in important_sentences[:max_iterations]:
            if count >= num_questions:
                break
            
            try:
                triggers = [" is ", " means ", " refers to ", " uses ", " makes ", " helps ", " allows ", " has ", " was ", " are "]
                trigger_found = None
                for t in triggers:
                    if t in sent:
                        trigger_found = t
                        break
                
                if not trigger_found:
                    continue

                parts = sent.split(trigger_found)
                if len(parts) < 2:
                    continue
                    
                subject = parts[0].strip()
                predicate = parts[1].strip()
                
                clean_subject = subject.strip(string.punctuation)
                if clean_subject.lower().startswith(("the ", "a ", "an ")):
                    clean_subject = clean_subject[4:]
                    
                if len(clean_subject.split()) > 3:
                    q_text = f"What {trigger_found.strip()} {subject}?"
                    ans = predicate
                else:
                    q_text = f"What {trigger_found.strip()} {predicate}?"
                    ans = clean_subject
                
                ans = ans.strip(string.punctuation)
                
                if not ans or len(ans) < 3 or len(ans.split()) > 4:
                    continue

                available_distractors = [w for w in smart_distractors if w.lower() != ans.lower()]
                random.shuffle(available_distractors)
                
                wrong_opts = available_distractors[:3]
                while len(wrong_opts) < 3:
                    wrong_opts.append("None of the above")
                
                options = [ans] + wrong_opts
                random.shuffle(options)
                
                opt_labels = ['A', 'B', 'C', 'D']
                formatted_opts = []
                correct_label = ""
                
                for i, opt in enumerate(options):
                    formatted_opts.append(f"{opt_labels[i]}) {opt}")
                    if opt == ans:
                        correct_label = opt_labels[i]
                
                mcqs.append({
                    "q": q_text,
                    "opts": formatted_opts,
                    "ans": correct_label
                })
                count += 1
                
            except Exception as e:
                # If one sentence fails, skip it and continue (Don't crash the whole app)
                safe_log(f"Skipping sentence due to logic error: {e}")
                continue
        
        return mcqs
        
    except Exception as e:
        safe_log(f"Generation Critical Error: {e}")
        return []

# --- SAFE FILE HANDLING ---

def read_file_safe(file):
    """Reads file with comprehensive error handling for public deployment."""
    if file is None:
        return None

    try:
        # 1. Check file size
        file_size = file.size
        if file_size == 0:
            st.error("⚠️ The uploaded file is empty.")
            return None
            
        if file_size > 5 * 1024 * 1024:  # 5MB
            st.error("⚠️ File too large. Please upload a file smaller than 5MB.")
            return None

        # 2. Handle Text Files
        if file.type == "text/plain":
            try:
                content = str(file.read(), "utf-8")
                if not content.strip():
                    st.error("⚠️ The text file is empty.")
                    return None
                return content
            except UnicodeDecodeError:
                # Fallback to latin-1
                content = str(file.read(), "latin-1")
                return content

        # 3. Handle PDF Files
        elif file.type == "application/pdf":
            if not pypdf:
                st.error("⚠️ PDF library not installed on server.")
                return None
            
            try:
                reader = pypdf.PdfReader(file)
                
                # Check if encrypted
                if reader.is_encrypted:
                    try:
                        # Try to decrypt with empty password (common for PDFs that look encrypted but aren't)
                        reader.decrypt("")
                    except Exception:
                        st.error("⚠️ PDF is password protected. Please remove password.")
                        return None
                    
                text = ""
                page_count = len(reader.pages)
                
                # Limit pages to prevent timeout on huge PDFs
                pages_to_read = min(page_count, 50) 
                
                for i in range(pages_to_read):
                    try:
                        extracted = reader.pages[i].extract_text()
                        if extracted:
                            text += extracted + " "
                    except Exception as page_err:
                        safe_log(f"Failed to read page {i}: {page_err}")
                        continue

                if not text.strip():
                    st.error("⚠️ Could not extract text from PDF. It might be an image-only PDF (scanned document).")
                    return None
                
                return text
                
            except PdfReadError:
                st.error("⚠️ Invalid or corrupted PDF file.")
                return None
            except DependencyError:
                st.error("⚠️ Missing system dependency for PDF parsing.")
                return None
            except Exception as e:
                safe_log(f"PDF Error: {e}")
                st.error("⚠️ An unknown error occurred while reading the PDF.")
                return None
        else:
            st.error("⚠️ Unsupported file type. Please use .txt or .pdf.")
            return None

    except Exception as e:
        safe_log(f"File Handling Unexpected Error: {e}")
        st.error("An unexpected error occurred while reading the file.")
        return None

# --- SAFE FILE CREATION ---

def create_pdf_safe(mcqs, title="Smart Exam Paper"):
    if not PDF_ENABLED:
        return None
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        pdf.cell(200, 10, txt=title, ln=1, align='C')
        pdf.ln(10)
        
        for i, item in enumerate(mcqs):
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, txt=f"Q{i+1}. {item['q']}", ln=1)
            
            pdf.set_font("Arial", size=11)
            for opt in item['opts']:
                pdf.cell(0, 10, txt=f"   {opt}", ln=1)
            
            pdf.ln(5)
        
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e:
        safe_log(f"PDF Gen Error: {e}")
        return None

def create_word_safe(mcqs):
    if not DOCX_ENABLED:
        return None
    try:
        doc = Document()
        doc.add_heading('Smart Exam Paper', 0)
        
        for i, item in enumerate(mcqs):
            p = doc.add_paragraph()
            p.add_run(f"Q{i+1}. {item['q']}").bold = True
            for opt in item['opts']:
                doc.add_paragraph(f"   {opt}")
            doc.add_paragraph()

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    except Exception as e:
        safe_log(f"Word Gen Error: {e}")
        return None

# --- MAIN APP ---

# Header
st.markdown('<div class="main-header">SP BROTHERS EXAM ENGINE</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Smart Exam Generator | AI-Powered Analysis</div>', unsafe_allow_html=True)

# Feature Status Check
missing_libs = []
if not PDF_ENABLED: missing_libs.append("fpdf")
if not DOCX_ENABLED: missing_libs.append("python-docx")

if missing_libs:
    st.warning(f"Note: Some features are disabled because libraries are missing: {', '.join(missing_libs)}. Please install them to enable full functionality.")

col1, col2 = st.columns([2, 1])
with col2:
    # Safety: Clamp max questions to prevent server load
    num_mcqs = st.number_input("Number of MCQs", min_value=1, max_value=100, value=10, step=5)
with col1:
    st.write("") 

st.markdown("### 📥 Upload Study Material")
input_method = st.radio("Method:", ("Paste Text", "Upload File"))

text_content = ""

if input_method == "Paste Text":
    text_content = st.text_area("Paste chapter notes or summary:", height=150, max_chars=50000)
else:
    file = st.file_uploader("Upload .txt or .pdf", type=['txt', 'pdf'])
    if file:
        with st.spinner("Reading file..."):
            text_content = read_file_safe(file)
            if text_content:
                st.success("File loaded successfully")

# Main Generation Logic with Global Try-Except
if st.button("GENERATE SMART EXAM", type="primary", use_container_width=True):
    try:
        if not text_content or len(text_content.strip()) < 50:
            st.warning("Please provide more text (at least 50 characters).")
        else:
            with st.spinner("Analyzing text..."):
                # Wrap generation in try/except to catch unhandled logic errors
                results = generate_smart_mcqs(text_content, num_mcqs)
                
                if not results:
                    st.warning("Could not generate questions. Try text with clearer definitions (e.g., 'X is Y').")
                else:
                    st.session_state.generated_mcqs = results
                    
                    if len(results) < num_mcqs:
                        st.info(f"⚠️ Only found {len(results)} 'important' sentences in the text.")
                        
    except Exception as e:
        safe_log(f"Main Button Unexpected Error: {e}")
        st.error("An unexpected system error occurred. Please try uploading different text.")

# Display Section
if st.session_state.generated_mcqs:
    st.markdown(f"### Exam Paper ({len(st.session_state.generated_mcqs)} Questions)")
    
    for i, item in enumerate(st.session_state.generated_mcqs):
        with st.container():
            st.markdown(f"<div class='mcq-container'>", unsafe_allow_html=True)
            st.markdown(f"<div class='question'>Q{i+1}. {item['q']}</div>", unsafe_allow_html=True)
            for opt in item['opts']:
                st.markdown(f"<div class='options'>{opt}</div>", unsafe_allow_html=True)
            
            if st.button("Show Answer", key=f"show_{i}"):
                st.markdown(f"<span class='answer-badge'>Correct Option: {item['ans']}</span>", unsafe_allow_html=True)
                
            st.markdown("</div>", unsafe_allow_html=True)

    # Downloads
    st.markdown("---")
    col_a, col_b, col_c = st.columns(3)
    
    key_content = "\n".join([f"Q{i+1}. {item['ans']}" for i, item in enumerate(st.session_state.generated_mcqs)])
    
    if PDF_ENABLED:
        pdf_data = create_pdf_safe(st.session_state.generated_mcqs)
        if pdf_data:
            col_a.download_button("📄 PDF", pdf_data, "Smart_Exam_Paper.pdf", "application/pdf")
        else:
            col_a.caption("PDF Generation Failed")
    else:
        col_a.write("PDF Disabled")

    if DOCX_ENABLED:
        word_data = create_word_safe(st.session_state.generated_mcqs)
        if word_data:
            col_b.download_button("📝 Word", word_data, "Smart_Exam_Paper.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            col_b.caption("Word Generation Failed")
    else:
        col_b.write("Word Disabled")

    col_c.download_button("🔑 Key (TXT)", key_content, "Answer_Key.txt", "text/plain")

# Footer
st.markdown("""
    <div class="footer-custom">
        <p>SP BROTHERS SOFTWARE SOLUTIONS | <span class="highlight">Developed by Muhammad Mudasar</span></p>
        <p style="margin-top: 5px; opacity: 0.6;">Smart Exam Generator | Ethical AI Use</p>
    </div>
""", unsafe_allow_html=True)